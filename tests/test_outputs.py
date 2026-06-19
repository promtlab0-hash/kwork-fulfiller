"""build_docx must render every shape the profiles emit, into a real .docx."""

import pytest

from engine import outputs

docx = pytest.importorskip("docx")  # skip cleanly if python-docx absent

NICHE = {"id": "t", "title": "Тестовый документ"}


def _text_of(path):
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_sections_shape(tmp_path):
    data = {"title": "Прайс", "sections": [
        {"heading": "Кофеварка", "level": 2, "paragraphs": ["Описание товара."],
         "bullets": ["пункт один", "пункт два"]},
    ]}
    files = outputs.build_docx(data, tmp_path, NICHE, "out")
    assert len(files) == 1 and files[0].exists() and files[0].stat().st_size > 0
    text = _text_of(files[0])
    assert "Кофеварка" in text and "пункт один" in text


def test_section_with_table(tmp_path):
    data = {"title": "КП", "sections": [
        {"heading": "Состав и цена", "paragraphs": ["Ниже смета."],
         "table": {"columns": ["Услуга", "Цена"], "rows": [["Аудит", "5000"], ["Настройка", "8000"]]}},
    ]}
    files = outputs.build_docx(data, tmp_path, NICHE, "kp")
    text = _text_of(files[0])
    assert "Услуга" in text and "Аудит" in text and "8000" in text


def test_flat_columns_rows_table(tmp_path):
    data = {"title": "Контент-план",
            "columns": ["Дата", "Тема", "Формат"],
            "rows": [["1 июня", "Запуск", "пост"], ["2 июня", "Кейс", "карусель"]]}
    files = outputs.build_docx(data, tmp_path, NICHE, "plan")
    text = _text_of(files[0])
    assert "Дата" in text and "карусель" in text


def test_flat_body_shape(tmp_path):
    data = {"title": "Док", "body": "## Раздел\n\nАбзац текста.\n\n- буллет"}
    files = outputs.build_docx(data, tmp_path, NICHE, "body")
    text = _text_of(files[0])
    assert "Раздел" in text and "буллет" in text


def test_add_table_accepts_dict_rows(tmp_path):
    # rows as dicts keyed by column — columns derived from keys when omitted.
    data = {"title": "T", "tables": [
        {"rows": [{"Имя": "A", "Балл": "10"}, {"Имя": "B", "Балл": "20"}]}]}
    files = outputs.build_docx(data, tmp_path, NICHE, "dict")
    text = _text_of(files[0])
    assert "Имя" in text and "Балл" in text and "20" in text


def test_unknown_builder_raises():
    with pytest.raises(KeyError):
        outputs.get_builder("does_not_exist")


def test_document_has_business_styling(tmp_path):
    from docx.shared import Cm, RGBColor
    data = {"title": "Отчёт", "sections": [{"heading": "Раздел", "paragraphs": ["Текст."]}]}
    files = outputs.build_docx(data, tmp_path, NICHE, "styled")
    document = docx.Document(str(files[0]))

    # Page margins are set (not Word defaults of 2.54cm/2.54cm).
    section = document.sections[0]
    assert abs(section.left_margin - Cm(2.5)) < Cm(0.1)
    assert abs(section.top_margin - Cm(2.5)) < Cm(0.1)

    # Footer carries a page-number label.
    footer_text = "\n".join(p.text for p in section.footer.paragraphs)
    assert "стр." in footer_text

    # Headings are neutral near-black, NOT a bright accent colour.
    heading = next(p for p in document.paragraphs if p.text == "Раздел")
    colour = heading.runs[0].font.color.rgb
    assert colour == RGBColor(0x1A, 0x1A, 0x1A)


def test_table_header_is_restrained_grey(tmp_path):
    from docx.oxml.ns import qn
    data = {"title": "T", "tables": [{"columns": ["A", "B"], "rows": [["1", "2"]]}]}
    files = outputs.build_docx(data, tmp_path, NICHE, "tbl")
    document = docx.Document(str(files[0]))
    header_cell = document.tables[0].rows[0].cells[0]
    shd = header_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == "595959"
