"""Builders — turn parsed LLM data into deliverable files.

Every builder takes (data, out_dir, niche, base_name) and returns the list of
pathlib.Path objects it created. Builders are pure with respect to `data`: they
never call the LLM.
"""

from __future__ import annotations

import csv
import html
import io
import json
import pathlib
import zipfile
from typing import Any

# python-docx / openpyxl are imported lazily inside builders so that builders
# which don't need them (html_pack, code_zip) still work if a dep is missing.


# --------------------------------------------------------------------------- #
# HTML pack (niche 1: prompts) — dark glass cards with one-click copy
# --------------------------------------------------------------------------- #

_HTML_CSS = """
:root{--bg:#05060c;--bg2:#0a0b14;--panel:rgba(20,22,34,.55);--panel-solid:#11121d;
--accent:#7c3aed;--accent-light:#b793f5;--accent-deep:#482289;--tint:#dac8fa;
--accent-rgb:124,58,237;--light-rgb:183,147,245;--deep-rgb:72,34,137;
--glow:rgba(var(--accent-rgb),.55);--line:rgba(var(--light-rgb),.22);
--line-strong:rgba(var(--light-rgb),.45);--txt:#eef0f7;--muted:#aeb4c8;--muted2:#7f879e}
*{box-sizing:border-box}html{scroll-behavior:smooth}
body{margin:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Arial,sans-serif;
color:var(--txt);background:radial-gradient(1200px 700px at 15% -10%,rgba(var(--deep-rgb),.40),transparent 60%),
radial-gradient(900px 600px at 110% 10%,rgba(var(--accent-rgb),.22),transparent 55%),
linear-gradient(180deg,var(--bg),var(--bg2));min-height:100vh;line-height:1.55}
.wrap{max-width:1080px;margin:0 auto;padding:0 18px}
header.hero{text-align:center;padding:60px 18px 24px}
.badge{display:inline-block;font-size:12px;letter-spacing:2.5px;text-transform:uppercase;
color:var(--accent-light);border:1px solid var(--line-strong);border-radius:100px;
padding:6px 16px;margin-bottom:20px;background:rgba(var(--accent-rgb),.08)}
.hero h1{font-size:clamp(28px,6vw,50px);margin:0 0 12px;font-weight:800;letter-spacing:-1px;
background:linear-gradient(180deg,#fff,var(--tint) 55%,var(--accent-light));
-webkit-background-clip:text;background-clip:text;color:transparent}
.hero p.sub{font-size:clamp(15px,2.4vw,18px);color:var(--muted);max-width:680px;margin:0 auto 8px}
.stats{display:flex;gap:14px;justify-content:center;flex-wrap:wrap;margin-top:24px}
.stat{background:var(--panel);border:1px solid var(--line);border-radius:16px;padding:14px 22px;min-width:110px}
.stat b{display:block;font-size:24px;color:#fff;line-height:1.1}.stat span{font-size:12px;color:var(--muted2)}
.searchbar{margin:26px 0 6px;position:relative}
.searchbar input{width:100%;background:var(--panel-solid);border:1px solid var(--line);
border-radius:14px;padding:15px 18px;color:var(--txt);font-size:15px;outline:none}
.searchbar input:focus{border-color:var(--line-strong)}
h2.cat{font-size:clamp(20px,3.6vw,27px);margin:42px 0 6px;display:flex;align-items:center;gap:12px;scroll-margin-top:20px}
h2.cat .ico{width:40px;height:40px;flex:0 0 40px;display:grid;place-items:center;border-radius:12px;
font-size:19px;background:linear-gradient(145deg,var(--accent-deep),var(--accent))}
.card{background:var(--panel);border:1px solid var(--line);border-radius:18px;padding:20px;margin:14px 0;transition:.18s}
.card:hover{border-color:var(--line-strong);transform:translateY(-2px)}
.card-head{display:flex;justify-content:space-between;align-items:flex-start;gap:14px;margin-bottom:8px}
.card-head h3{margin:0;font-size:17px;color:#fff;font-weight:700}
.card-num{font-size:12px;color:var(--muted2)}
.what{color:var(--muted);font-size:13.5px;margin:0 0 14px}.what b{color:var(--tint)}
pre.prompt{margin:0;white-space:pre-wrap;word-break:break-word;font-family:ui-monospace,Consolas,monospace;
font-size:13px;line-height:1.6;color:#dfe3f0;background:#0b0c15;
border:1px solid rgba(var(--light-rgb),.18);border-radius:12px;padding:16px;max-height:260px;overflow:auto}
.card-foot{display:flex;justify-content:flex-end;margin-top:13px}
.copy-btn{cursor:pointer;border:none;border-radius:11px;padding:11px 20px;font-size:14px;font-weight:600;
color:#fff;background:linear-gradient(145deg,var(--accent),var(--accent-deep));transition:.15s}
.copy-btn:hover{filter:brightness(1.12)}.copy-btn.done{background:linear-gradient(145deg,#10b981,#047857)}
footer{text-align:center;padding:46px 18px 56px;color:var(--muted2);font-size:13px;
border-top:1px solid var(--line);margin-top:46px}
footer .brand{color:var(--accent-light);font-weight:700;letter-spacing:1px}
#toast{position:fixed;left:50%;bottom:28px;transform:translateX(-50%) translateY(40px);
background:var(--panel-solid);border:1px solid var(--line-strong);color:#fff;padding:12px 22px;
border-radius:12px;font-size:14px;opacity:0;transition:.25s;pointer-events:none;z-index:99}
#toast.show{opacity:1;transform:translateX(-50%) translateY(0)}
""".strip()

_HTML_JS = """
const norm=s=>(s||'').toLowerCase();
function copyCard(btn,id){
  const pre=document.getElementById(id);
  navigator.clipboard.writeText(pre.innerText).then(()=>{
    btn.classList.add('done');btn.textContent='✓ Скопировано';
    showToast('Промпт скопирован в буфер');
    setTimeout(()=>{btn.classList.remove('done');btn.textContent='📋 Копировать';},1600);
  });
}
let toastTimer;
function showToast(msg){
  const t=document.getElementById('toast');t.textContent=msg;t.classList.add('show');
  clearTimeout(toastTimer);toastTimer=setTimeout(()=>t.classList.remove('show'),1800);
}
function filterCards(){
  const q=norm(document.getElementById('q').value);
  document.querySelectorAll('.card').forEach(c=>{
    c.style.display=norm(c.innerText).includes(q)?'':'none';
  });
}
""".strip()


def build_html_pack(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Render a prompt pack as a standalone HTML file plus a plain-text copy.

    Expected data shape:
      {"title": str, "subtitle": str,
       "categories": [{"name": str, "icon": str,
                       "prompts": [{"title", "prompt", "usage"}]}]}
    Also tolerates a flat {"prompts": [...]} list.
    """
    title = data.get("title") or niche.get("title", "Промпт-пак")
    subtitle = data.get("subtitle", "")
    categories = data.get("categories")
    if not categories:
        # Flat list -> single synthetic category.
        categories = [{"name": "Промпты", "icon": "✨", "prompts": data.get("prompts", [])}]

    total = sum(len(c.get("prompts", [])) for c in categories)

    cards_html: list[str] = []
    txt_lines: list[str] = [title, "=" * len(title), ""]
    idx = 0
    for cat in categories:
        cat_name = html.escape(cat.get("name", "Категория"))
        icon = html.escape(cat.get("icon", "•"))
        cards_html.append(
            f'<h2 class="cat"><span class="ico">{icon}</span>{cat_name}</h2>'
        )
        txt_lines.append(f"\n## {cat.get('name', 'Категория')}\n")
        for p in cat.get("prompts", []):
            idx += 1
            p_title = html.escape(p.get("title", f"Промпт {idx}"))
            p_text = html.escape(p.get("prompt", ""))
            usage = html.escape(p.get("usage", ""))
            pid = f"p{idx}"
            usage_html = f'<p class="what"><b>Когда:</b> {usage}</p>' if usage else ""
            cards_html.append(
                f'<div class="card"><div class="card-head">'
                f'<h3>{p_title}</h3><span class="card-num">№{idx}</span></div>'
                f"{usage_html}"
                f'<pre class="prompt" id="{pid}">{p_text}</pre>'
                f'<div class="card-foot"><button class="copy-btn" '
                f"onclick=\"copyCard(this,'{pid}')\">📋 Копировать</button></div></div>"
            )
            txt_lines.append(f"{idx}. {p.get('title', '')}")
            txt_lines.append(p.get("prompt", ""))
            if usage:
                txt_lines.append(f"Когда: {p.get('usage', '')}")
            txt_lines.append("")

    page = f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html.escape(title)}</title><style>{_HTML_CSS}</style></head><body>
<header class="hero"><div class="badge">kwork-fulfiller</div>
<h1>{html.escape(title)}</h1>
<p class="sub">{html.escape(subtitle)}</p>
<div class="stats"><div class="stat"><b>{total}</b><span>ПРОМПТОВ</span></div>
<div class="stat"><b>{len(categories)}</b><span>КАТЕГОРИЙ</span></div>
<div class="stat"><b>1 клик</b><span>КОПИРОВАНИЕ</span></div></div></header>
<div class="wrap">
<div class="searchbar"><input id="q" oninput="filterCards()" placeholder="🔎 Поиск по промптам…"></div>
{''.join(cards_html)}
</div>
<footer>Сгенерировано <span class="brand">kwork-fulfiller</span> • {total} промптов</footer>
<div id="toast"></div>
<script>{_HTML_JS}</script>
</body></html>"""

    html_path = out_dir / f"{base_name}.html"
    txt_path = out_dir / f"{base_name}.txt"
    html_path.write_text(page, encoding="utf-8")
    txt_path.write_text("\n".join(txt_lines).strip() + "\n", encoding="utf-8")
    return [html_path, txt_path]


# --------------------------------------------------------------------------- #
# Spreadsheets
# --------------------------------------------------------------------------- #


def _rows_from_data(data: Any, rows_key_candidates: list[str]) -> tuple[list[str], list[list[Any]]]:
    """Normalise data into (headers, rows). Accepts several shapes.

    {"columns": [...], "rows": [[...], ...]} -> used directly.
    {"<key>": [ {col: val}, ... ]}           -> headers from first dict.
    [ {col: val}, ... ]                      -> headers from first dict.
    """
    if isinstance(data, dict) and "columns" in data and "rows" in data:
        cols = [str(c) for c in data["columns"]]
        rows = [[r.get(c) if isinstance(r, dict) else r for c in cols]
                if isinstance(r, dict) else list(r) for r in data["rows"]]
        return cols, rows

    records: list[dict] | None = None
    if isinstance(data, list):
        records = [r for r in data if isinstance(r, dict)]
    elif isinstance(data, dict):
        for key in rows_key_candidates:
            val = data.get(key)
            if isinstance(val, list) and val and isinstance(val[0], dict):
                records = val
                break
        if records is None:
            # First list-of-dicts value, whatever its key.
            for val in data.values():
                if isinstance(val, list) and val and isinstance(val[0], dict):
                    records = val
                    break

    if not records:
        return ["value"], []

    headers: list[str] = []
    for rec in records:
        for k in rec:
            if k not in headers:
                headers.append(k)
    rows = [[rec.get(h, "") for h in headers] for rec in records]
    return headers, rows


def build_xlsx(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Write an .xlsx with a styled header row."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    rows_keys = ["clusters", "rows", "queries", "plan", "items", "cards", "posts", "data"]
    headers, rows = _rows_from_data(data, rows_keys)

    wb = Workbook()
    ws = wb.active
    ws.title = (niche.get("id") or "sheet")[:31]

    header_fill = PatternFill("solid", fgColor="4B2289")
    header_font = Font(bold=True, color="FFFFFF")
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center", wrap_text=True)

    for row in rows:
        ws.append([_cell(v) for v in row])

    for col_idx, header in enumerate(headers, start=1):
        max_len = len(str(header))
        for row in rows:
            if col_idx - 1 < len(row):
                max_len = max(max_len, len(str(row[col_idx - 1] or "")))
        letter = ws.cell(row=1, column=col_idx).column_letter
        ws.column_dimensions[letter].width = min(max(12, max_len + 2), 60)
    ws.freeze_panes = "A2"

    path = out_dir / f"{base_name}.xlsx"
    wb.save(path)
    return [path]


def build_csv(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Write a UTF-8 (BOM) .csv so Excel opens Cyrillic correctly."""
    rows_keys = ["clusters", "rows", "queries", "plan", "items", "cards", "posts", "data"]
    headers, rows = _rows_from_data(data, rows_keys)
    path = out_dir / f"{base_name}.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(headers)
        for row in rows:
            writer.writerow([_cell(v) for v in row])
    return [path]


def build_xlsx_csv(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Both spreadsheet formats — the typical SEO-core deliverable."""
    return build_xlsx(data, out_dir, niche, base_name) + build_csv(
        data, out_dir, niche, base_name
    )


def _cell(value: Any) -> Any:
    """Make a value safe for a spreadsheet cell."""
    if isinstance(value, (list, tuple)):
        return "\n".join(str(v) for v in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return value


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


def _add_docx_table(doc, table_spec: Any) -> None:
    """Render a {"columns": [...], "rows": [[...]]} spec as a styled Word table.

    Tolerant of shapes: rows may be lists (positional) or dicts (keyed by
    column). A purple header row matches the xlsx builder's palette. Silently
    ignores anything that isn't a usable table spec.
    """
    if not isinstance(table_spec, dict):
        return
    columns = [str(c) for c in (table_spec.get("columns") or [])]
    raw_rows = table_spec.get("rows") or []
    if not columns and raw_rows and isinstance(raw_rows[0], dict):
        # Derive columns from the first record's keys.
        for rec in raw_rows:
            if isinstance(rec, dict):
                for k in rec:
                    if k not in columns:
                        columns.append(k)
    if not columns:
        return

    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        cell = hdr[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(col)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Purple shading (4B2289) on the header cell.
        shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "4B2289",
        })
        cell._tc.get_or_add_tcPr().append(shd)

    for raw in raw_rows:
        if isinstance(raw, dict):
            values = [raw.get(c, "") for c in columns]
        elif isinstance(raw, (list, tuple)):
            values = list(raw)
        else:
            values = [raw]
        cells = table.add_row().cells
        for i in range(len(columns)):
            val = values[i] if i < len(values) else ""
            if isinstance(val, (list, tuple)):
                val = "\n".join(str(v) for v in val)
            cells[i].text = str(val)


def build_docx(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Render structured content as a Word document.

    Accepts either:
      {"title", "meta": {k: v}, "sections": [{"heading", "level", "paragraphs":
       [...], "bullets": [...], "table": {"columns", "rows"}}]}
    or a flat {"title", "body": "markdown-ish text"} which is split on blank
    lines and ## headings. A section may also carry a "table" spec, and the
    document may carry a top-level "tables": [spec, ...] appended after sections.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = data.get("title") or niche.get("title", "Документ")
    doc.add_heading(title, level=0)

    meta = data.get("meta") or {}
    if isinstance(meta, dict) and meta:
        for k, v in meta.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{k}: ")
            run.bold = True
            p.add_run(str(v))
        doc.add_paragraph()

    sections = data.get("sections")
    if sections:
        for sec in sections:
            heading = sec.get("heading")
            level = int(sec.get("level", 2))
            if heading:
                doc.add_heading(heading, level=max(1, min(level, 4)))
            for para in sec.get("paragraphs", []):
                doc.add_paragraph(str(para))
            for bullet in sec.get("bullets", []):
                doc.add_paragraph(str(bullet), style="List Bullet")
            if sec.get("table"):
                _add_docx_table(doc, sec["table"])
    elif data.get("tables") or data.get("table") or data.get("rows"):
        # Pure-table deliverable with no prose sections (e.g. content plan);
        # the actual table(s) are rendered by the top-level block below.
        pass
    else:
        body = data.get("body", "")
        if isinstance(body, list):
            body = "\n\n".join(str(b) for b in body)
        for block in str(body).split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=1)
            elif block.startswith(("- ", "* ")):
                for line in block.splitlines():
                    line = line.strip().lstrip("-* ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(block)

    # Top-level tables (a single spec via "table" or a list via "tables"),
    # rendered after any sections/body. Used by table-first deliverables.
    top_tables = data.get("tables")
    if isinstance(top_tables, dict):  # tolerate a single spec under "tables"
        top_tables = [top_tables]
    if not top_tables and data.get("table"):
        top_tables = [data["table"]]
    if not top_tables and isinstance(data, dict) and data.get("rows"):
        # Flat table deliverable: {"title", "columns": [...], "rows": [[...]]}.
        top_tables = [{"columns": data.get("columns"), "rows": data["rows"]}]
    for spec in (top_tables or []):
        _add_docx_table(doc, spec)

    # Make body text a touch larger for readability.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    path = out_dir / f"{base_name}.docx"
    doc.save(path)
    return [path]


def build_docx_xlsx(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Content-plan style deliverable: a table (.xlsx) + a prose doc (.docx).

    Expects {"plan": [...table rows...], "title", "sections"/"body"} — the plan
    feeds the spreadsheet, the rest feeds the doc.
    """
    files = build_xlsx(data, out_dir, niche, base_name)
    files += build_docx(data, out_dir, niche, base_name)
    return files


# --------------------------------------------------------------------------- #
# Code zip (niche 7: py_script)
# --------------------------------------------------------------------------- #


def build_code_zip(
    data: Any, out_dir: pathlib.Path, niche: dict, base_name: str
) -> list[pathlib.Path]:
    """Bundle script.py + README.md + requirements.txt into a .zip.

    Expected: {"script": "...py source...", "readme": "...", "requirements":
    "..." | [..], "title": str}. Also writes the loose files to out/ so the
    py_compiles check (which reads result["script"]) and a human can both see
    them.
    """
    script = data.get("script") or data.get("code") or ""
    readme = data.get("readme") or data.get("README") or f"# {data.get('title', 'Script')}\n"
    reqs = data.get("requirements") or data.get("requirements.txt") or ""
    if isinstance(reqs, (list, tuple)):
        reqs = "\n".join(str(r) for r in reqs)

    members = {
        "script.py": str(script),
        "README.md": str(readme),
        "requirements.txt": str(reqs),
    }

    zip_path = out_dir / f"{base_name}.zip"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in members.items():
            zf.writestr(name, content)
    zip_path.write_bytes(buf.getvalue())

    # Also drop loose copies for convenience / inspection.
    (out_dir / "script.py").write_text(str(script), encoding="utf-8")
    (out_dir / "README.md").write_text(str(readme), encoding="utf-8")
    (out_dir / "requirements.txt").write_text(str(reqs), encoding="utf-8")
    return [zip_path, out_dir / "script.py", out_dir / "README.md", out_dir / "requirements.txt"]


# --------------------------------------------------------------------------- #
# Registry
# --------------------------------------------------------------------------- #

BUILDERS = {
    "html_pack": build_html_pack,
    "xlsx": build_xlsx,
    "csv": build_csv,
    "xlsx_csv": build_xlsx_csv,
    "docx": build_docx,
    "docx_xlsx": build_docx_xlsx,
    "code_zip": build_code_zip,
}


def get_builder(name: str):
    """Look up a builder by name, raising a clear error if unknown."""
    try:
        return BUILDERS[name]
    except KeyError as exc:
        raise KeyError(
            f"Unknown builder {name!r}. Available: {', '.join(sorted(BUILDERS))}"
        ) from exc
